import os
import json
import fitz
import pathlib
from langchain.schema import Document
from langchain_openai import AzureChatOpenAI
from langchain_openai import AzureOpenAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader
from langchain_community.document_loaders import CSVLoader
from langchain_community.document_loaders.image import UnstructuredImageLoader
from langchain.chains import create_history_aware_retriever
from langchain_core.messages import AIMessage, HumanMessage
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
import uuid
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import PatternFill
import csv
import pandas as pd
from docx import Document as DocxDocument
from docx.enum.text import WD_COLOR_INDEX
from langchain.schema import Document
from typing import List
from langchain_community.retrievers import AzureAISearchRetriever
from langchain_community.document_loaders import AzureAIDocumentIntelligenceLoader
from vector_db import AzureSearch
import boto3
import shutil
from pathlib import Path
from openai import AzureOpenAI
import base64
import logging

logging.basicConfig(format="%(levelname)s     %(message)s", level=logging.INFO)
httpx_logger = logging.getLogger("httpx")
httpx_logger.setLevel(logging.WARNING)
logging.getLogger("uvicorn").setLevel(logging.WARNING)  # Set uvicorn to warning level
logging.getLogger("azure.core.pipeline.policies").setLevel(logging.WARNING)

async def upload_to_space(origin, output,remove, region_name='nyc3'):

    client = boto3.client(
        's3',
        region_name=region_name,
        endpoint_url=f'https://annotated-files.nyc3.digitaloceanspaces.com',
        aws_access_key_id=os.getenv("SPACES_ACCESS"),
        aws_secret_access_key=os.getenv("SPACES_SECRET")
    )
    
    try:
        client.upload_file(origin, "annotated-files", f"{output}", ExtraArgs={'ACL': 'public-read'})
        public_url = f'https://annotated-files.nyc3.digitaloceanspaces.com/annotated-files/{output}'

        if remove:
            os.remove(origin)
        return public_url
    
    except Exception as e:
        logging.info("error while placing file in bucket", e)
        return None


async def highlight_text_in_pdf(input_path, output_path, page_contents):
    doc = fitz.open(input_path)
    anotated_rects = []
    anotated_texts = []

    def get_line_rect(page, text_instance):
        block = page.get_text("dict")['blocks']
        for b in block:
            if b['type'] == 0:
                for line in b['lines']:
                    for span in line['spans']:
                        if (span['bbox'][0] <= text_instance.x0 <= span['bbox'][2] and
                            span['bbox'][1] <= text_instance.y0 <= span['bbox'][3]):
                            return fitz.Rect(span['bbox'][0], line['bbox'][1], span['bbox'][2], line['bbox'][3])
        return None

    try:
        for page_num, text_list in page_contents.items():
            if page_num >= len(doc) + 1 or page_num < 0:
                logging.info(f"Page number {page_num} is out of range.")
                continue

            page = doc[page_num - 1]
            for l in text_list:
                if l.strip():
                    text_instances = page.search_for(l)
                    if text_instances:
                        for inst in text_instances:
                            if (
                                inst not in anotated_rects
                            ):
                                line_rect = get_line_rect(page, inst)
                                if line_rect:
                                    annot = page.add_highlight_annot(line_rect)
                                    annot.update()
                                    anotated_rects.append(line_rect)
                                    anotated_texts.append(l)

        doc.save(output_path)
        doc.close()
        
    except Exception as e:
        logging.info("Error while marking PDF", e)


async def highlight_text_in_xlsx(input_path, output_path, page_contents):
    workbook = openpyxl.load_workbook(input_path)
    for page_num, details in page_contents.items():
        sheet_name = details['sheet']
        texts_to_highlight = details['text']

        if sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            row = sheet[page_num]
            for cell in row:
                for text in texts_to_highlight:
                    if text not in ["", "nan"] and text.strip() == (str(cell.value)).strip():
                        cell.fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")

    workbook.save(output_path)


async def highlight_text_in_csv(csv_file_path, xlsx_file_path, index_dict):
    wb = Workbook()
    ws = wb.active

    highlight_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")

    highlight_indices = set(index_dict.keys())
    # highlight_indices = [1,2,3,4,5]

    with open(csv_file_path, 'r', newline='') as csvfile:
        csvreader = csv.reader(csvfile)
        
        for index_csv, row in enumerate(csvreader):
            is_highlighted = index_csv in highlight_indices
            if is_highlighted:

                values = index_dict[index_csv]
                for col_val in row:
                    if any(val in col_val or col_val in val for val in values):
                        break

            ws.append(row)
            
            if is_highlighted:
                last_row = ws.max_row
                for cell in ws[last_row]:
                    cell.fill = highlight_fill

    wb.save(xlsx_file_path)

    logging.info(f"CSV file has been written to {xlsx_file_path}")


async def highlight_text_in_docx(docx_file, output_file, index_dict):
    doc = DocxDocument(docx_file)
    group_nums = sorted(index_dict.keys())  # Ensure keys are processed in order

    for para_index in group_nums:
        ending_index = para_index * 3 + 3
        starting_index = para_index * 3

        # Ensure the range of paragraphs does not exceed the document length
        for paragraph_index in range(starting_index, ending_index):
            if paragraph_index >= len(doc.paragraphs):  # Skip if out of range
                break

            paragraph = doc.paragraphs[paragraph_index]
            para_text = paragraph.text

            # Create a new paragraph to accumulate the highlighted text
            new_paragraph_runs = []

            # Split the paragraph text to highlight relevant parts
            pos = 0
            for q_text in index_dict.get(para_index, []):
                if q_text and q_text in para_text:
                    start_pos = para_text.find(q_text, pos)
                    end_pos = start_pos + len(q_text)
                    if start_pos != -1:
                        # Add text before the highlight
                        if start_pos > pos:
                            new_paragraph_runs.append((para_text[pos:start_pos], None))
                        # Add highlighted text
                        new_paragraph_runs.append((q_text, WD_COLOR_INDEX.YELLOW))
                        pos = end_pos

            # Add any remaining text after the last highlight
            if pos < len(para_text):
                new_paragraph_runs.append((para_text[pos:], None))

            # Clear the original paragraph and add the highlighted text
            paragraph.clear()
            for text, highlight in new_paragraph_runs:
                run = paragraph.add_run(text)
                if highlight:
                    run.font.highlight_color = highlight

    # Save the document with highlighted text
    doc.save(output_file)
    logging.info(f"Highlighted document saved as {output_file}")


async def docx_loader(file):
    docx = DocxDocument(file)
    documents_with_paragraphs = []
    current_group_content = []
    paragraph_group_number = 0

    for i, paragraph in enumerate(docx.paragraphs):
        # if paragraph.text.strip():  # Skip empty paragraphs
        current_group_content.append(paragraph.text)
        
        # Create a document for every 3 paragraphs
        if len(current_group_content) == 3:
            doc_with_group = Document(
                                        metadata={"source" : file, "page" : paragraph_group_number},
                                        id=str(uuid.uuid4()),
                                        page_content="\n".join(current_group_content)
                                        
                                    )

            documents_with_paragraphs.append(doc_with_group)
            current_group_content = []
            paragraph_group_number += 1
    
    # Handle remaining paragraphs
    if current_group_content:
        doc_with_group = Document(
                                    metadata={"source" : file, "page" : paragraph_group_number},
                                    id=str(uuid.uuid4()),
                                    page_content="\n".join(current_group_content)
                                )
        documents_with_paragraphs.append(doc_with_group)

    logging.info(f"Loaded documents from all paragraph groups: {len(documents_with_paragraphs)}")
    return documents_with_paragraphs


async def excel_loader(file):
    sheets = pd.read_excel(file, sheet_name=None)
    
    documents_with_rows = []
    
    for sheet_name, df in sheets.items():
        for i, row in df.iterrows():
            row_text = "\n".join(str(cell) for cell in row)
            doc_with_row = Document(
                metadata={
                    "sheet": sheet_name,   
                    "row": i + 1,          
                    "source": file         
                },
                page_content=row_text      
            )
            documents_with_rows.append(doc_with_row)
    
    logging.info(f"Loaded documents from all sheets with row and column numbers: {len(documents_with_rows)}")
    return documents_with_rows

async def add_source(documents, file):
    
    for doc in documents:
        if "source" not in doc.metadata:
            doc.metadata["source"] = file

    return documents

    
failed_files = []

async def load_data(filenames):
    logging.info("Background task initiated")
    try:
        all_documents = []

        files = os.path.join(os.getcwd(), "docs")
        for filename in filenames:
            try:
                file = os.path.abspath(os.path.join(str(files), str(filename)))
                logging.info(f"Processing {file}")
                file_extension = pathlib.Path(file).suffix

                if file_extension.lower() == ".pdf":
                    try:
                        raw_documents = PyPDFLoader(file, extract_images=True).load()
                    except ValueError as e:
                        logging.info(f"Failed to extract images from {file}: {e}")
                        raw_documents = PyPDFLoader(file, extract_images=False).load()
                    all_documents.extend(raw_documents)

                elif file_extension.lower() == ".xlsx":
                    logging.info("Loading")
                    raw_documents = await excel_loader(file)
                    all_documents.extend(raw_documents)

                elif file_extension.lower() == ".csv":
                    raw_documents = CSVLoader(file_path=file).load()
                    all_documents.extend(raw_documents)

                elif file_extension.lower() == ".docx":
                    raw_documents = await docx_loader(file)
                    all_documents.extend(raw_documents)

                elif file_extension.lower() == ".txt":
                    raw_documents = TextLoader(file).load()
                    all_documents.extend(raw_documents)

                elif file_extension.lower() in [".jpg", ".jpeg", ".png"]:                    
                    raw_documents = AzureAIDocumentIntelligenceLoader(
                        api_endpoint=os.getenv("IMAGE_LOADER_ENDPOINT"), api_key=os.getenv("IMAGE_LOADER_KEY"), file_path=file, mode="page",
                    ).load()
                    raw_documents = await add_source(raw_documents, file)
                    all_documents.extend(raw_documents)

            except Exception as e:
                logging.info(f"Failed to process {filename}: {e}")
                failed_files.append(os.path.basename(filename))
                return 

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=300, chunk_overlap=50
        )
        texts = text_splitter.split_documents(all_documents)

        logging.info("split")
        embedding = AzureOpenAIEmbeddings(
            model="text-embedding-ada-002",
            azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT_EMBEDDINGS"),
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT_EMBEDDINGS"),
            api_key=os.getenv("AZURE_OPENAI_API_KEY_EMBEDDINGS"),
        )  

        logging.info("embeddings fetched")
        vectordb = AzureSearch(
                azure_search_endpoint=os.getenv("AZURE_SEARCH_ENDPOINT"),
                azure_search_key=os.getenv("AZURE_SEARCH_KEY"),
                index_name="soc-index",  # Replace with your index name
                embedding_function=embedding.embed_query,
            )
        
        logging.info("db fetched")

        logging.info("embeddings created")

        for text in texts:

            if "id" not in text:
                text.id = str(uuid.uuid4())

            text.metadata["source"] = text.metadata["source"].split("/")[-1]


            if "row" in text.metadata:
                text.metadata["page"] = text.metadata['row']
                del text.metadata["row"]

            if "sheet" not in text.metadata:
                text.metadata["sheet"] = ""

        logging.info(texts[0])
        await vectordb.aadd_documents(documents=texts)

        logging.info("Files Added")

        return filenames

    except Exception as e:
        logging.info(f"Error in load_data: {e}")
        return False


async def check_documents_exist(source):
    source = source.replace(" ", "_")

    if source in failed_files:
        return False, "Failed Processing"
    
    retriever = AzureAISearchRetriever(
        api_key=os.getenv("AZURE_SEARCH_KEY"),
        service_name="azure-vector-db",
        index_name="soc-index",
        top_k=1,
        filter=f"metadata/source eq '{source}'"
    )

    documents = retriever.invoke("")
    if len(documents) > 0:
        doc_source = documents[0].metadata['metadata']['source']
        if doc_source == source:
            logging.info(f"Doc Exists")
            return True, "File Exists"
        
    return False, "File Processing"

chat_history = {}


fallback_texts = ["Your question is not relevant to the evidence",
                  'Try phrasing your question to be more specific to the evidence',
                  'Your question is not relevant to the evidence',
                  "Your question is not relevant to the evidence. Try phrasing your question to be more specific to the evidence.",
                  "Hi", "Hello", "provided documents do not contain", "provided documents do not include"
                  ]


async def clean_content(response, source):
    
    client = AzureOpenAI(
        azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT"), 
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),  
        api_version="2024-02-15-preview",
        azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
        )

    user_question = f"Question: {response['input']}, or answer: {response['answer']}?"
    system_prompt = """
        Your task is to filter irrelevant content based on the provided question or answer:
        Question & Answer: {user_question}.
        Please return only the contexts that are relevant to this question or answer.

        
        If the Question or Answer is a general talk, such as greetings, small talk (e.g., "hi", "hello", "how are you", "what's the weather", "how's it going", etc.), 
        then the "context" should be an empty list.
        
        Maintain the format of the context as the original!

        Respond in similar JSON format.
        "context" : [...], This is a list of dicts
    """
    system_prompt = system_prompt.format(user_question=user_question, source=source, answer=response['answer'])
    context_message = f"Here is the context to filter:\n{response['context']}"
    try:
        response_ai = client.chat.completions.create(
            model="gpt-4o", 
            response_format={"type": "json_object"},
            # temperature=0.3,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": context_message}
            ]
        )
        response_text = response_ai.choices[0].message.content.strip()
        filtered_response = json.loads(response_text)
        logging.info(f"filtered response {filtered_response}")
        response['context'] = filtered_response['context']
        # if filtered_response["answer"] != response["answer"]:
        #     response["answer"] = filtered_response["answer"]

        return response

    except Exception as e:
        logging.info(f"Passing context as is due to Error: {e}")
        return response


async def check_file_format(persist_directory: str):
    # Mapping of file extensions to output values
    file_format_output = {
        ".pdf": (5, 7),
        ".csv": (5, 7),
        ".docx": (5, 7),
        ".xlsx": (5, 7)
    }

    # Extract the file extension and return the corresponding value
    file_extension = Path(persist_directory).suffix.lower()
    return file_format_output.get(file_extension, (1,5))


async def create_chain(retriever, model):
    system_prompt = '''You are an expert Software Security Auditor. 
    Your job is to provide answers relevant to the knowledge base provided.  
    Do not provide any information that is not contained in the documents retrieved.
    Always give summarized answers using only the content from the retrieved documents. 
    Greetings are allowed but Do not answer any other general queries, always stay within the scope of the documents. 
    If a query is not within the scope of the documents say 'Your question is not relevant to the evidence'
    {context}
    '''
    
    main_prompt = ChatPromptTemplate.from_messages(
        [
            ("system", system_prompt),
            MessagesPlaceholder(variable_name="chat_history"),
            ("human", "{input}"),
        ]
    )

    retriever_prompt = ChatPromptTemplate.from_messages(
        [
            MessagesPlaceholder(variable_name="chat_history"),
            ("human", "{input}"),
            (
                "human",
                "Always Fetch and use the documents provided for any information and use those to formulate a response",
            ),
        ]
    )

    chain = create_stuff_documents_chain(llm=model, prompt=main_prompt)

    # No need to create a separate retriever here; using AzureAISearchRetriever directly
    # history_aware_retriever = create_history_aware_retriever(
    #     llm=model, retriever=retriever, prompt=retriever_prompt
    # )
    # retrieval_chain = create_retrieval_chain(retriever, chain)


    return create_retrieval_chain(retriever, chain)


async def process_chat(chain, question, chat_history, dir, threshold):
    # Invoke the chain with input question and chat history
    response = chain.invoke({"input": question, "chat_history": chat_history})
    # response = chain.invoke({"input": question})
    
    answer = response['answer']

    
    final_response ={
        "input" : question,
        "chat_history" : chat_history,
        "answer": answer,
        "context": []
    }
    logging.info(f"first response  {final_response}")

    for docs in response["context"]:
        score = docs.metadata['@search.score']
        metadata_dict = docs.metadata["metadata"]
        if metadata_dict['source'] == dir:
            if score >= threshold:
                custom_data = {"metadata" : metadata_dict, "page_content" : docs.page_content}
                final_response['context'].append(custom_data)
        else:
            final_response["answer"] = "Your question is not relevant to the evidence"

    logging.info(f"initial response {final_response}")
    return final_response


async def generate_response(uid, persist_directory, rfe, markup):

    persist_directory = persist_directory.replace(" ", "_")
    
    chat_history.setdefault(uid, [])

    threshold, k = await check_file_format(persist_directory)

    try:
        retriever = AzureAISearchRetriever(
            api_key=os.getenv("AZURE_SEARCH_KEY"),
            service_name="azure-vector-db",
            index_name="soc-index",
            top_k=k,  # Number of documents to retrieve
            filter=f"metadata/source eq '{persist_directory}'"
        )
        # Initialize Azure Chat model
        model = AzureChatOpenAI(
            max_tokens=200,
            openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
            azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
            api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        )
        # Create chain with Azure Cognitive Search retriever and model
        chain = await create_chain(retriever, model)

        # Process chat with the created chain
        result = await process_chat(chain, rfe, chat_history[uid], persist_directory, threshold)

        # logging.info(f"original answer {result['answer']}")
        # logging.info(f"original context {result['context']}")
        
        result = await clean_content(result, persist_directory)

        # logging.info(f"filtered answer {result['answer']}")
        # logging.info(f"filtered context {result['context']}")
        
        ai_answer = result["answer"].strip()

        chat_history[uid]= [HumanMessage(content=rfe), AIMessage(content=result["answer"])]

        for text in fallback_texts:
            if text in ai_answer or text in result["answer"]:
                return {
                "AI_message": ai_answer,
                "Source": None,
                "Pages/Rows" : None,
                "Annotated_file" : None
                }
        

        source = persist_directory
        pages, page_contents = set(), {}
        markup_check = False
        for doc_details in result["context"]:
            
            if "page_content" in doc_details:
                markup_check = True

                lines = doc_details["page_content"].splitlines()

            
                if "pdf" in source:
                    page = int(doc_details['metadata'].get('page')) + 1
                    page_contents[page] = lines

                elif "page_name" in source:
                    page = doc_details['metadata'].get('page_name')
                    page_contents[page] = lines

                elif source.endswith((".xlsx", ".csv")):
                    page = int(doc_details['metadata'].get('page')) + 1
                    page_contents[page] = {"sheet": doc_details['metadata'].get('sheet'), "text" : lines}

                elif "docx" in source:
                    page = int(doc_details['metadata'].get('page'))
                    page_contents[page] = doc_details["page_content"].split("\n")
                else:
                    page = 0        

            pages.add(page)
        
        space_url = ""

        if markup and markup_check:
            if "pdf" in source:
                await highlight_text_in_pdf(
                                            f"./docs/{source}",
                                            "out.pdf",
                                            page_contents,
                                            )    

                space_file_path = f"annotated_{source}"
                space_url = await upload_to_space("out.pdf", space_file_path, True)

            elif "xlsx" in source:
                await highlight_text_in_xlsx(
                                            f"./docs/{source}",
                                            "out.xlsx", 
                                            page_contents
                                            )
                space_file_path = f"annotated_{source}"
                space_url = await upload_to_space("out.xlsx", space_file_path, True)

            elif "csv" in source:
                await highlight_text_in_csv(
                                            f"./docs/{source}",
                                            "out.xlsx",
                                            page_contents
                                            )
                space_file_path = f"annotated_{source}"
                space_file_path = space_file_path.replace("csv", "xlsx")
                space_url = await upload_to_space("out.xlsx", space_file_path, True)

            elif "docx" in source:
                await highlight_text_in_docx(
                                            f"./docs/{source}",
                                            "out.docx",
                                            page_contents
                                            )
                space_file_path = f"annotated_{source}"
                space_url = await upload_to_space("out.docx", space_file_path, True)

        return {
            "AI_message": ai_answer,
            "Source": source,
            "Pages/Rows" : pages,
            "Annotated_file" : space_url
            }
    
    except Exception as e:
        logging.info(f"Error occured {e}")

        return {
            "AI_message": "There was an issue while fetching information",
            "Source": "",
            "Pages/Rows" : "",
            "Annotated_file" : ""
            }